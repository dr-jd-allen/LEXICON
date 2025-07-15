#!/usr/bin/env python3
"""
Document processor for LEXICON MVP
Converts PDFs to clean text using unstructured library
"""

import os
from pathlib import Path
from typing import List, Dict
from unstructured.partition.pdf import partition_pdf
from unstructured.chunking.basic import chunk_elements
import json


class DocumentProcessor:
    def __init__(self, input_dir: str = "data/raw", output_dir: str = "data/processed"):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def process_pdf(self, pdf_path: Path) -> Dict:
        """Process a single PDF and return structured content."""
        print(f"Processing {pdf_path.name}...")
        
        # Extract elements from PDF
        elements = partition_pdf(
            filename=str(pdf_path),
            strategy="hi_res",  # High resolution for better accuracy
            extract_images_in_pdf=False,  # Skip images for text focus
            infer_table_structure=True  # Preserve table structures
        )
        
        # Chunk elements for better RAG performance
        chunks = chunk_elements(
            elements,
            max_characters=1000,
            overlap=100
        )
        
        # Structure the output
        relative_path = pdf_path.relative_to(self.input_dir) if pdf_path.is_relative_to(self.input_dir) else pdf_path
        processed_doc = {
            "filename": pdf_path.name,
            "relative_path": str(relative_path),
            "directory": str(relative_path.parent),
            "doc_type": self._classify_document(pdf_path),
            "chunks": []
        }
        
        for i, chunk in enumerate(chunks):
            processed_doc["chunks"].append({
                "chunk_id": f"{pdf_path.stem}_chunk_{i}",
                "text": str(chunk),
                "metadata": {
                    "page_number": getattr(chunk, "page_number", None),
                    "element_type": chunk.category
                }
            })
            
        return processed_doc
    
    def process_text_file(self, txt_path: Path) -> Dict:
        """Process a text file and return structured content."""
        print(f"Processing text file {txt_path.name}...")
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Chunk the text
        chunks = []
        chunk_size = 1000
        overlap = 100
        
        for i in range(0, len(text), chunk_size - overlap):
            chunk_text = text[i:i + chunk_size]
            if chunk_text.strip():
                chunks.append(chunk_text)
        
        # Structure the output
        relative_path = txt_path.relative_to(self.input_dir) if txt_path.is_relative_to(self.input_dir) else txt_path
        processed_doc = {
            "filename": txt_path.name,
            "relative_path": str(relative_path),
            "directory": str(relative_path.parent),
            "doc_type": self._classify_document(txt_path),
            "chunks": []
        }
        
        for i, chunk in enumerate(chunks):
            processed_doc["chunks"].append({
                "chunk_id": f"{txt_path.stem}_chunk_{i}",
                "text": chunk,
                "metadata": {
                    "page_number": None,
                    "element_type": "text"
                }
            })
            
        return processed_doc
    
    def process_docx_file(self, docx_path: Path) -> Dict:
        """Process a .docx file and return structured content."""
        print(f"Processing DOCX file {docx_path.name}...")
        
        try:
            from docx import Document
            doc = Document(str(docx_path))
            
            # Extract all text
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            text = "\n\n".join(full_text)
            
            # Chunk the text
            chunks = []
            chunk_size = 1000
            overlap = 100
            
            for i in range(0, len(text), chunk_size - overlap):
                chunk_text = text[i:i + chunk_size]
                if chunk_text.strip():
                    chunks.append(chunk_text)
            
            # Structure the output
            relative_path = docx_path.relative_to(self.input_dir) if docx_path.is_relative_to(self.input_dir) else docx_path
            processed_doc = {
                "filename": docx_path.name,
                "relative_path": str(relative_path),
                "directory": str(relative_path.parent),
                "doc_type": self._classify_document(docx_path),
                "chunks": []
            }
            
            for i, chunk in enumerate(chunks):
                processed_doc["chunks"].append({
                    "chunk_id": f"{docx_path.stem}_chunk_{i}",
                    "text": chunk,
                    "metadata": {
                        "page_number": None,
                        "element_type": "text"
                    }
                })
                
            return processed_doc
            
        except ImportError:
            print("python-docx not installed. Processing as binary file.")
            # Fallback to treating it as text
            return self.process_text_file(docx_path)
    
    def _classify_document(self, pdf_path: Path) -> str:
        """Classify document type based on filename and directory structure."""
        filename_lower = pdf_path.name.lower()
        path_str = str(pdf_path).lower()
        
        # Check directory structure first for more accurate classification
        if "defense expert reports" in path_str or "plaintiff expert reports" in path_str:
            return "expert_report"
        elif "defense-cited tbi research" in path_str or "relevant medical literature" in path_str:
            return "medical_literature"
        elif "deposition" in filename_lower or "depo" in filename_lower:
            return "deposition"
        elif "medical records" in path_str or "neuro-psych assessment" in path_str:
            return "medical_record"
        elif "favorable rulings" in path_str or "motions in limine" in path_str:
            return "legal_precedent"
        elif "defense counsel strategies" in path_str or "plaintiff counsel strategies" in path_str:
            return "legal_strategy"
        # Fallback to filename patterns
        elif "expert" in filename_lower and ("report" in filename_lower or "cv" in filename_lower):
            return "expert_report"
        elif "medical" in filename_lower or "record" in filename_lower:
            return "medical_record"
        elif "daubert" in filename_lower or "motion" in filename_lower or "mil" in filename_lower:
            return "legal_motion"
        else:
            return "other"
    
    def process_batch(self, files: List[str] = None) -> None:
        """Process multiple PDFs and text files in batch."""
        if files is None:
            # Process all PDFs and .txt files in input directory and subdirectories
            pdf_files = list(self.input_dir.rglob("*.pdf"))
            txt_files = list(self.input_dir.rglob("*.txt"))
            # Also look for .docx files we can process
            docx_files = list(self.input_dir.rglob("*.docx"))
            # Exclude temporary docx files
            docx_files = [f for f in docx_files if not f.name.startswith("~")]
            
            all_files = pdf_files + txt_files + docx_files
            print(f"Found {len(pdf_files)} PDF files, {len(txt_files)} TXT files, and {len(docx_files)} DOCX files")
            print(f"Total files to process: {len(all_files)}")
        else:
            all_files = [self.input_dir / f for f in files]
            
        results = []
        
        for file_path in all_files:
            if file_path.exists():
                try:
                    # Process based on file type
                    if file_path.suffix.lower() == '.pdf':
                        processed = self.process_pdf(file_path)
                    elif file_path.suffix.lower() == '.txt':
                        processed = self.process_text_file(file_path)
                    elif file_path.suffix.lower() == '.docx':
                        processed = self.process_docx_file(file_path)
                    else:
                        continue
                    
                    # Create output directory structure matching input
                    relative_path = file_path.relative_to(self.input_dir) if file_path.is_relative_to(self.input_dir) else file_path.name
                    output_subdir = self.output_dir / relative_path.parent
                    output_subdir.mkdir(parents=True, exist_ok=True)
                    
                    # Save as JSON for Dify ingestion
                    output_path = output_subdir / f"{file_path.stem}.json"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(processed, f, indent=2, ensure_ascii=False)
                        
                    # Also save as plain text for simpler use cases
                    text_output = output_subdir / f"{file_path.stem}.txt"
                    full_text = "\n\n".join([chunk["text"] for chunk in processed["chunks"]])
                    with open(text_output, 'w', encoding='utf-8') as f:
                        f.write(full_text)
                        
                    results.append({
                        "file": file_path.name,
                        "relative_path": str(relative_path),
                        "doc_type": processed["doc_type"],
                        "status": "success",
                        "chunks": len(processed["chunks"])
                    })
                    
                except Exception as e:
                    print(f"Error processing {file_path.name}: {str(e)}")
                    results.append({
                        "file": file_path.name,
                        "status": "error",
                        "error": str(e)
                    })
            else:
                print(f"File not found: {file_path}")
                
        # Save processing summary
        summary_path = self.output_dir / "processing_summary.json"
        with open(summary_path, 'w') as f:
            json.dump(results, f, indent=2)
            
        print(f"\nProcessing complete. Summary saved to {summary_path}")
        

    def get_documents_by_type(self, doc_type: str = None) -> List[Path]:
        """Get all processed documents of a specific type."""
        json_files = list(self.output_dir.rglob("*.json"))
        documents = []
        
        for json_file in json_files:
            with open(json_file, 'r') as f:
                doc_data = json.load(f)
                if doc_type is None or doc_data.get('doc_type') == doc_type:
                    documents.append({
                        'path': json_file,
                        'doc_type': doc_data.get('doc_type'),
                        'relative_path': doc_data.get('relative_path'),
                        'filename': doc_data.get('filename')
                    })
        
        return documents
    
    def generate_processing_report(self) -> Dict:
        """Generate a summary report of all processed documents by category."""
        json_files = list(self.output_dir.rglob("*.json"))
        report = {
            'total_documents': len(json_files),
            'by_type': {},
            'by_directory': {}
        }
        
        for json_file in json_files:
            with open(json_file, 'r') as f:
                doc_data = json.load(f)
                doc_type = doc_data.get('doc_type', 'unknown')
                directory = doc_data.get('directory', 'root')
                
                # Count by type
                if doc_type not in report['by_type']:
                    report['by_type'][doc_type] = []
                report['by_type'][doc_type].append(doc_data.get('filename'))
                
                # Count by directory
                if directory not in report['by_directory']:
                    report['by_directory'][directory] = []
                report['by_directory'][directory].append(doc_data.get('filename'))
        
        # Save report
        report_path = self.output_dir / "document_processing_report.json"
        with open(report_path, 'w') as f:
            json.dump(report, f, indent=2)
            
        return report


if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    
    # Process all PDFs in the data/raw directory structure
    processor.process_batch()
    
    # Generate summary report
    report = processor.generate_processing_report()
    print(f"\nProcessing Report:")
    print(f"Total documents processed: {report['total_documents']}")
    print(f"\nDocuments by type:")
    for doc_type, files in report['by_type'].items():
        print(f"  {doc_type}: {len(files)} documents")