#!/usr/bin/env python3
"""
DOCX generation for LEXICON MVP
Creates formatted Daubert motions from LLM output
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path


class DaubertMotionGenerator:
    def __init__(self, template_path: str = "templates/daubert_template.docx"):
        self.template_path = Path(template_path)
        
    def create_template(self):
        """Create a basic Daubert motion template if it doesn't exist."""
        if self.template_path.exists():
            return
            
        self.template_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # Add custom styles
        styles = doc.styles
        
        # Case caption style
        caption_style = styles.add_style('Case Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = 'Times New Roman'
        caption_style.font.size = Pt(12)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Motion title style
        title_style = styles.add_style('Motion Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Add template content
        doc.add_paragraph('IN THE UNITED STATES DISTRICT COURT', style='Case Caption')
        doc.add_paragraph('FOR THE [DISTRICT]', style='Case Caption')
        doc.add_paragraph()
        
        # Case caption table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = '[PLAINTIFF NAME],'
        table.cell(1, 0).text = 'Plaintiff,'
        table.cell(2, 0).text = 'v.'
        table.cell(0, 1).text = 'Case No. [CASE NUMBER]'
        table.cell(1, 1).text = ''
        table.cell(2, 1).text = '[DEFENDANT NAME],'
        
        doc.add_paragraph()
        doc.add_paragraph('MOTION TO EXCLUDE EXPERT TESTIMONY UNDER DAUBERT', style='Motion Title')
        
        # Introduction
        doc.add_heading('INTRODUCTION', level=1)
        doc.add_paragraph(
            'Defendant respectfully moves this Court to exclude the testimony of '
            '[EXPERT NAME] under Federal Rule of Evidence 702 and Daubert v. '
            'Merrell Dow Pharmaceuticals, Inc., 509 U.S. 579 (1993).'
        )
        
        # Legal Standard
        doc.add_heading('LEGAL STANDARD', level=1)
        doc.add_paragraph(
            'Under Rule 702 and Daubert, expert testimony is admissible only if it is '
            'both relevant and reliable. The proponent of expert testimony must establish '
            'by a preponderance of the evidence that the testimony meets these requirements. '
            'Daubert, 509 U.S. at 592-93. The Court serves as a "gatekeeper" to ensure that '
            'speculative, unreliable expert testimony does not reach the jury. '
            'Kumho Tire Co. v. Carmichael, 526 U.S. 137, 147 (1999).'
        )
        
        doc.add_paragraph(
            'In assessing reliability, courts consider several factors including: '
            '(1) whether the theory or technique can be tested; '
            '(2) whether it has been subjected to peer review and publication; '
            '(3) the known or potential rate of error; '
            '(4) the existence of standards controlling the technique\'s operation; and '
            '(5) general acceptance in the relevant scientific community. '
            'Daubert, 509 U.S. at 593-94.'
        )
        
        # Argument section placeholder
        doc.add_heading('ARGUMENT', level=1)
        doc.add_paragraph('[ARGUMENT_PLACEHOLDER]')
        
        # Conclusion
        doc.add_heading('CONCLUSION', level=1)
        doc.add_paragraph(
            'For the foregoing reasons, Defendant respectfully requests that this Court '
            'exclude the testimony of [EXPERT NAME] in its entirety.'
        )
        
        doc.add_paragraph()
        doc.add_paragraph('Respectfully submitted,')
        doc.add_paragraph()
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('[ATTORNEY NAME]')
        doc.add_paragraph('[BAR NUMBER]')
        doc.add_paragraph('[LAW FIRM]')
        doc.add_paragraph('[ADDRESS]')
        doc.add_paragraph('[PHONE]')
        doc.add_paragraph('[EMAIL]')
        doc.add_paragraph('Attorney for Defendant')
        
        # Save template
        doc.save(str(self.template_path))
        print(f"Template created at {self.template_path}")
        
    def generate_motion(self, 
                       argument_text: str,
                       case_info: Dict[str, str],
                       output_filename: str = None) -> Path:
        """Generate a formatted Daubert motion from the LLM output."""
        
        # Load template
        if not self.template_path.exists():
            self.create_template()
            
        doc = Document(str(self.template_path))
        
        # Replace placeholders
        replacements = {
            '[DISTRICT]': case_info.get('district', '[DISTRICT]'),
            '[PLAINTIFF NAME]': case_info.get('plaintiff', '[PLAINTIFF NAME]'),
            '[DEFENDANT NAME]': case_info.get('defendant', '[DEFENDANT NAME]'),
            '[CASE NUMBER]': case_info.get('case_number', '[CASE NUMBER]'),
            '[EXPERT NAME]': case_info.get('expert_name', '[EXPERT NAME]'),
            '[ATTORNEY NAME]': case_info.get('attorney_name', '[ATTORNEY NAME]'),
            '[BAR NUMBER]': case_info.get('bar_number', '[BAR NUMBER]'),
            '[LAW FIRM]': case_info.get('law_firm', '[LAW FIRM]'),
            '[ADDRESS]': case_info.get('address', '[ADDRESS]'),
            '[PHONE]': case_info.get('phone', '[PHONE]'),
            '[EMAIL]': case_info.get('email', '[EMAIL]'),
            '[ARGUMENT_PLACEHOLDER]': argument_text
        }
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
        
        # Generate output filename
        if output_filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"Daubert_Motion_{case_info.get('expert_name', 'Expert').replace(' ', '_')}_{timestamp}.docx"
            
        output_path = Path('output') / output_filename
        output_path.parent.mkdir(exist_ok=True)
        
        # Save document
        doc.save(str(output_path))
        print(f"Motion generated: {output_path}")
        
        return output_path


if __name__ == "__main__":
    # Test template creation
    generator = DaubertMotionGenerator()
    generator.create_template()
    
    # Test motion generation
    test_case_info = {
        'district': 'NORTHERN DISTRICT OF CALIFORNIA',
        'plaintiff': 'John Doe',
        'defendant': 'ABC Corporation',
        'case_number': '3:24-cv-12345',
        'expert_name': 'Dr. Jane Smith',
        'attorney_name': 'Jane Attorney',
        'bar_number': '123456',
        'law_firm': 'Smith & Associates',
        'address': '123 Main St, Suite 100, San Francisco, CA 94105',
        'phone': '(415) 555-1234',
        'email': 'jattorney@smithlaw.com'
    }
    
    test_argument = "This is a test argument section for the Daubert motion."
    
    generator.generate_motion(test_argument, test_case_info)